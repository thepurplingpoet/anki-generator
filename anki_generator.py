#!/usr/bin/env python3
"""
Anki Deck Generator
===================
Reads PDF, DOCX, or TXT files, analyzes them with an LLM (Claude API or
local Ollama), and generates Anki-importable CSV decks (or pushes via AnkiConnect).

Requires: pip install pdfplumber pypdf python-docx
Optional: brew install poppler (for better PDF image extraction)
          brew install ollama  (for local LLM backend)

Usage:
    python anki_generator.py <input_file> [options]

Backends (pick one):
    (default)         Claude API (needs ANTHROPIC_API_KEY)
    --ollama          Local LLM via Ollama (free, private, no API key)
    --offline         Rule-based heuristics (no LLM at all)

Options:
    --output, -o      Output CSV filename
    --deck, -d        Anki deck name (default: derived from filename)
    --model           Ollama model name (default: llama3.1:8b)
    --max-cards, -m   Target number of cards to generate
    --cards-per-page  Target cards per page (e.g. 5)
    --pages, -p       Page range for PDFs (e.g. "1-50")
    --low-mem         Low memory mode for 8GB Macs
    --anki-connect    Push cards directly to Anki via AnkiConnect
    --api-key         Anthropic API key (or set ANTHROPIC_API_KEY env var)
    --verbose, -v     Show detailed progress
"""

import argparse
import csv
import json
import os
import re
import sys
import hashlib
import base64
import socket
from pathlib import Path
from typing import Optional


def _check_poppler_installed() -> bool:
    """Check if poppler-utils (pdfimages) is available."""
    import subprocess
    try:
        subprocess.run(["pdfimages", "-v"], capture_output=True, timeout=5)
        return True
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        return False


POPPLER_AVAILABLE = _check_poppler_installed()

# ─── Caption Extraction ─────────────────────────────────────────────────────

# Patterns for image captions/labels in books and textbooks
# Matches: "Figure 1-1: Description", "Fig. 3.2 — Some caption", "Diagram 4:", 
#          "FIGURE 1-1 Single server", "Table 2.3: Data", etc.
CAPTION_PATTERNS = [
    # "Figure 1-1: Description text" or "Figure 1-1. Description text"  
    r'((?:Fig(?:ure)?|Diagram|Illustration|Chart|Graph|Image|Photo|Exhibit|Plate|Map|Schematic|Table)\s*\.?\s*(\d+[\.\-]\d*[a-z]?))\s*[:\.\-—–]\s*(.+)',
    # "FIGURE 1-1 Description" (no separator, ALL CAPS label)
    r'((?:FIGURE|DIAGRAM|TABLE|CHART|GRAPH|IMAGE|ILLUSTRATION)\s+(\d+[\.\-]\d*[a-z]?))\s+([A-Z].+)',
    # "Figure 4a: Description" (number + letter, no separator in label)
    r'((?:Fig(?:ure)?|Diagram|Table|Chart|Graph|Image)\s*\.?\s*(\d+[a-z]))\s*[:\.\-—–]\s*(.+)',
    # "Fig 3 — Description" (single number)
    r'((?:Fig(?:ure)?|Diagram|Table|Chart)\s*\.?\s*(\d+))\s*[:\.\-—–]\s*(.+)',
]

_caption_compiled = [re.compile(p, re.IGNORECASE) for p in CAPTION_PATTERNS]


def _extract_caption(page_text: str, image_index: int = 0) -> tuple[str, str]:
    """Extract figure caption and label from page text.
    
    Args:
        page_text: The text content of the page
        image_index: Which image on the page (0-based) to find the caption for
    
    Returns:
        (full_caption, label) e.g. ("Figure 1-1: Single server setup", "1-1")
        Returns ("", "") if no caption found.
    """
    captions_found = []
    
    for line in page_text.split('\n'):
        line = line.strip()
        if not line:
            continue
        for pattern in _caption_compiled:
            m = pattern.match(line)
            if m:
                full_label_text = m.group(1).strip()   # "Figure 1-1"
                label_number = m.group(2).strip()       # "1-1"
                description = m.group(3).strip()        # "Single server setup"
                full_caption = f"{full_label_text}: {description}"
                captions_found.append((full_caption, label_number))
                break
    
    if captions_found:
        # Return the caption matching the image_index (if multiple on one page)
        idx = min(image_index, len(captions_found) - 1)
        return captions_found[idx]
    
    return ("", "")


def _build_caption_index(images: list[dict]) -> dict:
    """Build a mapping from caption labels to image names.
    
    Returns dict like:
        {"1-1": "img_3_0.png", "1.2": "img_4_0.png", "2-1": "img_8_0.png"}
    Also includes normalized versions: "1-1" → "1.1" and vice versa.
    """
    label_to_image = {}
    for img in images:
        label = img.get("caption_label", "")
        if label:
            label_to_image[label] = img["name"]
            # Also store normalized versions (1-1 ↔ 1.1)
            normalized = label.replace('-', '.').replace('_', '.')
            label_to_image[normalized] = img["name"]
            normalized2 = label.replace('.', '-').replace('_', '-')
            label_to_image[normalized2] = img["name"]
    return label_to_image


# ─── File Readers ───────────────────────────────────────────────────────────

def read_pdf(filepath: str, page_range: tuple = None) -> tuple[str, list[dict]]:
    """Extract text and images from a PDF file.
    
    Args:
        filepath: Path to PDF
        page_range: Optional (start, end) tuple, 1-indexed inclusive. 
                    e.g. (1, 20) reads pages 1 through 20.
    """
    import pdfplumber
    from pypdf import PdfReader

    text_parts = []
    images = []

    # Extract text with pdfplumber (better layout preservation)
    with pdfplumber.open(filepath) as pdf:
        total_pages = len(pdf.pages)
        
        # Determine which pages to read
        if page_range:
            start = max(0, page_range[0] - 1)  # convert to 0-indexed
            end = min(total_pages, page_range[1])
            pages_to_read = list(range(start, end))
            print(f"  Reading pages {page_range[0]}-{page_range[1]} of {total_pages}")
        else:
            pages_to_read = list(range(total_pages))
            print(f"  Reading all {total_pages} pages")
        
        for i in pages_to_read:
            page = pdf.pages[i]
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"[Page {i+1}]\n{page_text}")

            # Extract tables as formatted text
            tables = page.extract_tables()
            for j, table in enumerate(tables):
                if table:
                    table_text = format_table(table)
                    text_parts.append(f"[Page {i+1} - Table {j+1}]\n{table_text}")

    # Build page text lookup for caption extraction
    page_text_map = {}
    for part in text_parts:
        m = re.match(r'\[Page\s+(\d+)\]\n(.*)', part, re.DOTALL)
        if m:
            pg = int(m.group(1))
            page_text_map[pg] = m.group(2)

    # Extract images using multiple methods (most reliable first)
    first_page = page_range[0] if page_range else 1
    last_page = page_range[1] if page_range else None
    
    images = _extract_images_from_pdf(filepath, page_text_map, first_page, last_page)

    full_text = "\n\n".join(text_parts)
    return full_text, images


def _extract_images_from_pdf(filepath: str, page_text_map: dict,
                              first_page: int = 1, last_page: int = None) -> list[dict]:
    """Extract images from PDF using multiple methods for maximum reliability.
    
    Methods tried in order:
    1. pdfimages CLI (poppler-utils) — most reliable, extracts ALL embedded images
    2. pdfplumber — good for many PDFs
    3. pypdf — fallback
    """
    import subprocess
    import tempfile
    import glob

    images = []

    # Method 1: pdfimages (from poppler-utils) — extracts images per page
    try:
        with tempfile.TemporaryDirectory() as tmpdir:
            cmd = ["pdfimages", "-png"]
            if first_page:
                cmd += ["-f", str(first_page)]
            if last_page:
                cmd += ["-l", str(last_page)]
            cmd += [filepath, os.path.join(tmpdir, "img")]
            
            result = subprocess.run(cmd, capture_output=True, timeout=120)
            
            if result.returncode == 0:
                # pdfimages outputs: img-000.png, img-001.png, etc.
                img_files = sorted(glob.glob(os.path.join(tmpdir, "img-*.png")))
                
                if img_files:
                    # pdfimages doesn't tell us which page each image is from directly
                    # Use -list mode to get page numbers
                    list_cmd = ["pdfimages", "-list"]
                    if first_page:
                        list_cmd += ["-f", str(first_page)]
                    if last_page:
                        list_cmd += ["-l", str(last_page)]
                    list_cmd += [filepath]
                    
                    list_result = subprocess.run(list_cmd, capture_output=True, text=True, timeout=30)
                    
                    # Parse the list output to get page numbers
                    # Format: "page   num  type   width  height  ..."
                    page_numbers = []
                    if list_result.returncode == 0:
                        for line in list_result.stdout.strip().split('\n')[2:]:  # skip header
                            parts = line.split()
                            if parts and parts[0].isdigit():
                                page_numbers.append(int(parts[0]))
                    
                    for idx, img_file in enumerate(img_files):
                        with open(img_file, 'rb') as f:
                            img_data = f.read()
                        
                        # Skip tiny images (icons, bullets, etc.) — less than 5KB
                        if len(img_data) < 5000:
                            continue
                        
                        page_num = page_numbers[idx] if idx < len(page_numbers) else first_page + idx
                        img_name = f"img_{page_num}_{idx}.png"
                        
                        # Get caption from page text
                        caption, caption_label = "", ""
                        pg_text = page_text_map.get(page_num, "")
                        if pg_text:
                            # Count how many images we've already assigned to this page
                            same_page_count = sum(1 for i in images if i.get("page") == page_num)
                            caption, caption_label = _extract_caption(pg_text, same_page_count)
                        
                        images.append({
                            "name": img_name,
                            "data": base64.b64encode(img_data).decode('utf-8'),
                            "page": page_num,
                            "caption": caption,
                            "caption_label": caption_label,
                            "description": caption or f"Image from page {page_num}"
                        })
                    
                    if images:
                        print(f"  Extracted {len(images)} images via pdfimages")
                        return images

    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        # pdfimages not installed — try next method
        pass

    # Method 2: pdfplumber image extraction
    try:
        import pdfplumber
        with pdfplumber.open(filepath) as pdf:
            img_count = 0
            page_start = (first_page - 1) if first_page else 0
            page_end = last_page if last_page else len(pdf.pages)
            
            for i in range(page_start, min(page_end, len(pdf.pages))):
                page = pdf.pages[i]
                page_num = i + 1
                
                if hasattr(page, 'images') and page.images:
                    page_img_idx = 0
                    for img_info in page.images:
                        try:
                            # pdfplumber gives image bbox but we need the actual bytes
                            # Extract via the page's crop
                            x0 = img_info.get('x0', 0)
                            y0 = img_info.get('top', 0)
                            x1 = img_info.get('x1', page.width)
                            y1 = img_info.get('bottom', page.height)
                            
                            cropped = page.crop((x0, y0, x1, y1))
                            pil_img = cropped.to_image(resolution=150).original
                            
                            import io
                            buf = io.BytesIO()
                            pil_img.save(buf, format='PNG')
                            img_data = buf.getvalue()
                            
                            if len(img_data) < 5000:
                                continue
                            
                            img_name = f"img_{page_num}_{img_count}.png"
                            caption, caption_label = "", ""
                            pg_text = page_text_map.get(page_num, "")
                            if pg_text:
                                caption, caption_label = _extract_caption(pg_text, page_img_idx)
                            
                            images.append({
                                "name": img_name,
                                "data": base64.b64encode(img_data).decode('utf-8'),
                                "page": page_num,
                                "caption": caption,
                                "caption_label": caption_label,
                                "description": caption or f"Image from page {page_num}"
                            })
                            img_count += 1
                            page_img_idx += 1
                        except Exception:
                            pass
            
            if images:
                print(f"  Extracted {len(images)} images via pdfplumber")
                return images
    except Exception:
        pass

    # Method 3: pypdf fallback (original method — least reliable)
    try:
        from pypdf import PdfReader
        reader = PdfReader(filepath)
        img_count = 0
        page_start = (first_page - 1) if first_page else 0
        page_end = last_page if last_page else len(reader.pages)
        
        for page_num_idx in range(page_start, min(page_end, len(reader.pages))):
            page = reader.pages[page_num_idx]
            page_num = page_num_idx + 1
            if hasattr(page, 'images'):
                page_img_idx = 0
                for img in page.images:
                    try:
                        img_data = img.data
                        if len(img_data) < 5000:
                            continue
                        img_name = f"img_{page_num}_{img_count}.png"
                        caption, caption_label = "", ""
                        pg_text = page_text_map.get(page_num, "")
                        if pg_text:
                            caption, caption_label = _extract_caption(pg_text, page_img_idx)
                        
                        images.append({
                            "name": img_name,
                            "data": base64.b64encode(img_data).decode('utf-8'),
                            "page": page_num,
                            "caption": caption,
                            "caption_label": caption_label,
                            "description": caption or f"Image from page {page_num}"
                        })
                        img_count += 1
                        page_img_idx += 1
                    except Exception:
                        pass
        
        if images:
            print(f"  Extracted {len(images)} images via pypdf")
    except Exception:
        pass

    if not images:
        print("  No images could be extracted. Install poppler for best results:")
        print("    macOS:  brew install poppler")
        print("    Linux:  sudo apt install poppler-utils")

    return images


def read_docx(filepath: str) -> tuple[str, list[dict]]:
    """Extract text and images from a DOCX file."""
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    doc = Document(filepath)
    text_parts = []
    images = []
    img_count = 0

    for para in doc.paragraphs:
        if para.text.strip():
            # Preserve heading structure
            if para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                text_parts.append(f"{'#' * int(level) if level.isdigit() else '#'} {para.text}")
            else:
                text_parts.append(para.text)

    # Extract tables
    for i, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
        if rows:
            table_text = format_table(rows)
            text_parts.append(f"[Table {i+1}]\n{table_text}")

    # Extract images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                ext = rel.target_ref.split('.')[-1] if '.' in rel.target_ref else 'png'
                img_name = f"img_{img_count}.{ext}"
                images.append({
                    "name": img_name,
                    "data": base64.b64encode(img_data).decode('utf-8'),
                    "page": 0,
                    "description": f"Embedded image {img_count + 1}"
                })
                img_count += 1
            except Exception:
                pass

    full_text = "\n\n".join(text_parts)
    return full_text, images


def read_txt(filepath: str) -> tuple[str, list[dict]]:
    """Read a plain text file."""
    with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
        text = f.read()
    return text, []


def format_table(table: list[list]) -> str:
    """Format a table as readable text."""
    if not table:
        return ""
    lines = []
    for row in table:
        cleaned = [str(cell).strip() if cell else "" for cell in row]
        lines.append(" | ".join(cleaned))
    return "\n".join(lines)


# ─── Content Filter (skip boilerplate) ─────────────────────────────────────

# Section headings that are NOT study material
SKIP_SECTIONS = {
    # Front/back matter
    "contents", "table of contents", "toc",
    "introduction", "preface", "foreword", "prologue",
    "acknowledgements", "acknowledgments", "acknowledgement",
    "dedication", "about the author", "about the authors", "author biography",
    "contributors", "about this book",
    # End matter
    "index", "subject index", "author index", "name index", "keyword index",
    "bibliography", "references", "works cited", "further reading",
    "suggested reading", "recommended reading", "selected bibliography",
    "appendix", "appendices",
    "glossary",  # glossary is borderline — but usually duplicates definitions in text
    "endnotes", "footnotes", "notes",
    "credits", "permissions", "copyright",
    "colophon", "publishing information", "cataloging-in-publication",
    "isbn", "edition notice",
    # Common textbook filler
    "learning objectives", "key terms list", "chapter outline",
    "review questions", "discussion questions", "practice problems",
    "answers to exercises", "answer key", "solutions manual",
    "list of figures", "list of tables", "list of illustrations",
    "list of abbreviations", "list of symbols", "nomenclature",
}

# Patterns for page-level boilerplate to strip out
SKIP_PAGE_PATTERNS = [
    r'^\s*©.*$',                          # Copyright lines
    r'^\s*All rights reserved.*$',
    r'^\s*ISBN[\s:\-]*[\dX\-]+',          # ISBN lines
    r'^\s*Printed in\b.*$',
    r'^\s*Published by\b.*$',
    r'^\s*Library of Congress\b.*$',
    r'^\s*First (published|edition|printing)\b.*$',
    r'^\s*\d+\s*$',                       # Standalone page numbers
    r'^\s*(Chapter\s+\d+\s*$)',           # Bare "Chapter 5" lines (no title)
]

_skip_page_re = [re.compile(p, re.IGNORECASE | re.MULTILINE) for p in SKIP_PAGE_PATTERNS]


def is_skip_section(heading: str) -> bool:
    """Check if a heading indicates a non-study section to skip."""
    cleaned = heading.strip().lower()
    # Remove numbering like "1.", "Chapter 1:", "A.", "I."
    cleaned = re.sub(r'^(chapter\s+)?\d+[\.\):\s]+', '', cleaned)
    cleaned = re.sub(r'^[a-z][\.\):\s]+', '', cleaned)
    cleaned = re.sub(r'^[ivxlcdm]+[\.\):\s]+', '', cleaned, flags=re.IGNORECASE)
    cleaned = cleaned.strip().rstrip('.')

    return cleaned in SKIP_SECTIONS


def filter_content(text: str) -> str:
    """Remove boilerplate sections and non-study content from extracted text.
    
    Detects section headings (markdown # or [Page] markers) and skips
    everything under a boilerplate heading until the next heading.
    """
    lines = text.split('\n')
    filtered_lines = []
    skipping = False
    skip_reason = ""
    skipped_sections = []

    for line in lines:
        # Detect section headings: "# Heading", "[Page N] HEADING", or ALL-CAPS lines
        heading = None

        # Markdown heading
        h_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if h_match:
            heading = h_match.group(2)

        # Page marker sometimes followed by section title
        page_match = re.match(r'^\[Page\s+\d+\]\s*$', line)
        if page_match:
            # Page marker alone — check next significant line (handled by flow)
            pass

        # ALL CAPS heading (common in PDFs): at least 3 words, all uppercase
        if not heading and line.strip():
            stripped = line.strip()
            if (stripped.isupper() and len(stripped.split()) >= 1
                    and len(stripped) < 100 and not stripped.startswith('|')):
                heading = stripped

        # Check if this heading triggers a skip
        if heading and is_skip_section(heading):
            if not skipping:
                skipped_sections.append(heading.strip())
            skipping = True
            skip_reason = heading.strip()
            continue

        # A new non-skip heading ends the skip
        if heading and not is_skip_section(heading):
            skipping = False

        if not skipping:
            # Clean page-level boilerplate from individual lines
            should_skip_line = False
            for pattern in _skip_page_re:
                if pattern.search(line):
                    should_skip_line = True
                    break
            if not should_skip_line:
                filtered_lines.append(line)

    if skipped_sections:
        unique = list(dict.fromkeys(skipped_sections))  # preserve order, dedupe
        print(f"  Skipped non-study sections: {', '.join(unique[:8])}")
        if len(unique) > 8:
            print(f"    ... and {len(unique) - 8} more")

    result = '\n'.join(filtered_lines)

    # Clean up excessive blank lines left by filtering
    result = re.sub(r'\n{4,}', '\n\n\n', result)

    return result


# ─── Card Generation Prompt ────────────────────────────────────────────────

SYSTEM_PROMPT = """You are an expert at creating Anki flashcards for effective learning and retention.
You will be given the full text of a document. Your job is to create flashcards that:

1. Cover ALL important concepts, facts, definitions, relationships, and details in the material.
2. Use the Basic note type: each card has a Front and a Back.
3. The BACK must contain EXACT content from the original text. You may add supplementary context,
   but the core answer must use the original wording verbatim. If there is a clear typo or
   grammatical error in the original, fix it and note "[corrected]".
4. The FRONT should test recall effectively. Choose the best type for each card:
   - "question": A clear question that tests a specific piece of knowledge
   - "cloze": A sentence from the text with a key term blanked out as {{c1::answer}}
   - "image": If an image is critical to understanding (reference it by name)
5. Assign topic TAGS to each card (lowercase, hyphens for spaces, hierarchical with ::).
   Example tags: "biology::cell-structure", "history::world-war-2::causes"
6. Create granular cards — one concept per card. Avoid mega-cards.
7. For lists/enumerations, create both individual cards AND a card testing the full list.
8. For definitions, create a card asking "What is X?" and possibly a reverse card.
9. For cause-effect relationships, test both directions.

Return your response as a JSON array of objects with this schema:
{
  "cards": [
    {
      "front_type": "question" | "cloze" | "image",
      "front": "string — the question, cloze text, or image reference + question",
      "back": "string — the answer using EXACT text from the source",
      "tags": ["tag1", "tag2"],
      "source_quote": "string — the exact sentence(s) from the source this card is based on"
    }
  ]
}

IMPORTANT: Return ONLY valid JSON. No markdown fences, no commentary outside the JSON."""

# Compact prompt for small models (phi3:mini, etc.) — saves ~300 tokens of context
SYSTEM_PROMPT_COMPACT = """Create Anki flashcards from the text below. Return JSON only.

Rules:
- Create MANY cards (at least 3-5 per paragraph). Cover every fact, definition, and concept.
- Back: use EXACT words from the source text
- Front types: "question" (ask about the fact) or "cloze" (blank out a key term with {{c1::term}})
- Add topic tags

JSON format:
{"cards":[{"front_type":"question","front":"What is X?","back":"exact text from source","tags":["topic"],"source_quote":"exact text"}]}

IMPORTANT: Generate AS MANY cards as possible. Every sentence with a fact should become a card."""

def build_user_prompt(text: str, has_images: bool, image_names: list[str] = None,
                      max_cards: int = None, min_cards: int = None) -> str:
    prompt = f"""Analyze the following document and create comprehensive Anki flashcards.

DOCUMENT TEXT:
---
{text}
---
"""
    if has_images and image_names:
        prompt += f"""
The document contains these images: {', '.join(image_names)}
If any image is essential for understanding a concept, create an "image" type card referencing it.
"""
    
    if max_cards and min_cards:
        prompt += f"\nYou MUST generate between {min_cards} and {max_cards} cards. Aim for {max_cards}."
    elif max_cards:
        prompt += f"\nGenerate at least {max(max_cards - 5, 3)} cards, up to {max_cards} cards."
    elif min_cards:
        prompt += f"\nYou MUST generate at least {min_cards} cards. More is better."
    
    prompt += """
Generate flashcards covering ALL important material. Return ONLY valid JSON."""
    return prompt


# ─── Claude API Integration ────────────────────────────────────────────────

def call_claude_api(text: str, has_images: bool, image_names: list[str] = None,
                    api_key: str = None, max_cards: int = None) -> list[dict]:
    """Call Claude API to generate flashcards from text."""
    import urllib.request
    import urllib.error

    if not api_key:
        api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    if not api_key:
        print("ERROR: No Anthropic API key found.")
        print("Set ANTHROPIC_API_KEY environment variable or pass --api-key")
        sys.exit(1)

    user_prompt = build_user_prompt(text, has_images, image_names, max_cards=max_cards)

    # Chunk very long texts to stay within context limits
    # Claude Sonnet handles ~200k tokens; we'll be conservative
    MAX_TEXT_CHARS = 150_000
    if len(text) > MAX_TEXT_CHARS:
        print(f"  Document is very long ({len(text)} chars). Processing in chunks...")
        return _process_in_chunks(text, has_images, image_names, api_key, max_cards)

    payload = {
        "model": "claude-sonnet-4-20250514",
        "max_tokens": 8192,
        "system": SYSTEM_PROMPT,
        "messages": [
            {"role": "user", "content": user_prompt}
        ]
    }

    data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(
        "https://api.anthropic.com/v1/messages",
        data=data,
        headers={
            "Content-Type": "application/json",
            "x-api-key": api_key,
            "anthropic-version": "2023-06-01"
        }
    )

    print("  Calling Claude API to generate flashcards...")
    try:
        with urllib.request.urlopen(req, timeout=120) as resp:
            result = json.loads(resp.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else "No details"
        print(f"  API Error {e.code}: {error_body}")
        sys.exit(1)

    # Extract text content
    response_text = ""
    for block in result.get("content", []):
        if block.get("type") == "text":
            response_text += block["text"]

    # Parse JSON
    return _parse_cards_json(response_text)


def _process_in_chunks(text: str, has_images: bool, image_names: list[str],
                       api_key: str, max_cards: int) -> list[dict]:
    """Process very long documents in chunks."""
    CHUNK_SIZE = 120_000
    chunks = []
    for i in range(0, len(text), CHUNK_SIZE):
        chunks.append(text[i:i + CHUNK_SIZE])

    all_cards = []
    cards_per_chunk = (max_cards // len(chunks)) + 1 if max_cards else None

    for i, chunk in enumerate(chunks):
        print(f"  Processing chunk {i+1}/{len(chunks)}...")
        cards = call_claude_api(chunk, has_images, image_names, api_key, cards_per_chunk)
        all_cards.extend(cards)

    return all_cards


def _repair_json(text: str) -> str:
    """Attempt to repair common JSON issues from LLM output."""
    # 1. Fix newlines inside string values (most common issue)
    #    Strategy: replace literal newlines inside quoted strings with \\n
    repaired = []
    in_string = False
    escape_next = False
    for ch in text:
        if escape_next:
            repaired.append(ch)
            escape_next = False
            continue
        if ch == '\\':
            repaired.append(ch)
            escape_next = True
            continue
        if ch == '"' and not escape_next:
            in_string = not in_string
            repaired.append(ch)
            continue
        if in_string and ch == '\n':
            repaired.append('\\n')
            continue
        if in_string and ch == '\r':
            continue
        if in_string and ch == '\t':
            repaired.append('\\t')
            continue
        repaired.append(ch)
    text = ''.join(repaired)

    # 2. Fix trailing commas before } or ]
    text = re.sub(r',\s*([}\]])', r'\1', text)

    # 3. Fix single quotes used as string delimiters (risky but common)
    #    Only do this if there are no double quotes at all
    if '"' not in text and "'" in text:
        text = text.replace("'", '"')

    # 4. Fix truncated JSON — try to close open brackets/braces
    open_braces = text.count('{') - text.count('}')
    open_brackets = text.count('[') - text.count(']')
    
    if open_braces > 0 or open_brackets > 0:
        # Truncated mid-output. Find the last cleanly closed card object.
        # A complete card ends with "}" possibly followed by ","
        # We want to keep everything up to and including the last complete "}"
        # that is part of a card (not a nested structure like tags array)
        
        # Find all positions of "}," which indicate a complete card followed by another
        last_card_boundary = -1
        for m in re.finditer(r'\}\s*,', text):
            last_card_boundary = m.end()
        
        if last_card_boundary > 0:
            text = text[:last_card_boundary - 1]  # remove the trailing comma
            # Strip any trailing comma
            text = text.rstrip().rstrip(',')
        else:
            # No complete card boundary found; try last complete }
            last_brace = text.rfind('}')
            if last_brace > 0:
                text = text[:last_brace + 1]

        # Re-count and close remaining structures
        open_braces = text.count('{') - text.count('}')
        open_brackets = text.count('[') - text.count(']')
    
    # Close any remaining open structures (in correct order: ] then })
    text += ']' * max(0, open_brackets)
    text += '}' * max(0, open_braces)

    return text


def _parse_cards_json(text: str) -> list[dict]:
    """Parse the JSON response from the LLM, handling common issues."""
    # Strip markdown fences if present
    text = text.strip()
    text = re.sub(r'^```json\s*', '', text)
    text = re.sub(r'\s*```$', '', text)
    text = text.strip()

    # Try parsing as-is first
    parsed = None
    try:
        parsed = json.loads(text)
    except json.JSONDecodeError:
        pass

    # Try repairing common issues
    if parsed is None:
        try:
            repaired = _repair_json(text)
            parsed = json.loads(repaired)
            print("  (repaired malformed JSON from LLM)")
        except json.JSONDecodeError:
            pass

    # Try extracting JSON object from surrounding text
    if parsed is None:
        match = re.search(r'\{[\s\S]*\}', text)
        if match:
            try:
                parsed = json.loads(_repair_json(match.group()))
            except json.JSONDecodeError:
                pass

    # Try extracting JSON array
    if parsed is None:
        match = re.search(r'\[[\s\S]*\]', text)
        if match:
            try:
                parsed = json.loads(_repair_json(match.group()))
            except json.JSONDecodeError:
                pass

    if parsed is None:
        print(f"  ERROR: Could not parse API response as JSON even after repair.")
        print(f"  Response preview: {text[:500]}")
        return []

    # Extract the cards list from various possible structures
    cards_raw = []
    if isinstance(parsed, dict):
        # Try common wrapper keys: "cards", "flashcards", "notes", "deck", "data"
        for key in ["cards", "flashcards", "notes", "deck", "data", "items", "questions"]:
            if key in parsed and isinstance(parsed[key], list):
                cards_raw = parsed[key]
                break
        if not cards_raw:
            # Maybe the dict itself is a single card
            if any(k in parsed for k in ["front", "question", "prompt", "q"]):
                cards_raw = [parsed]
            else:
                # Try the first list value in the dict
                for v in parsed.values():
                    if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                        cards_raw = v
                        break
    elif isinstance(parsed, list):
        cards_raw = parsed

    if not cards_raw:
        print("  WARNING: Could not find cards in the response.")
        print(f"  Response structure: {type(parsed).__name__}")
        if isinstance(parsed, dict):
            print(f"  Keys: {list(parsed.keys())}")
        return []

    # Normalize each card to our expected schema
    normalized = []
    for raw in cards_raw:
        if not isinstance(raw, dict):
            continue
        card = _normalize_card(raw)
        if card:
            normalized.append(card)

    if len(normalized) < len(cards_raw):
        skipped = len(cards_raw) - len(normalized)
        print(f"  WARNING: {skipped} malformed cards skipped during normalization")

    return normalized


def _normalize_card(raw: dict) -> dict:
    """Normalize a card dict from various LLM output formats to our standard schema.
    
    Handles key variations like:
      front/question/prompt/q → front
      back/answer/response/a → back
      front_type/type/card_type/format → front_type
      tags/topics/categories/labels → tags
      source_quote/source/reference/quote → source_quote
    """
    # ── Front ──
    front = None
    for key in ["front", "question", "prompt", "q", "front_text", "card_front",
                "front_side", "clue", "stimulus"]:
        if key in raw and raw[key]:
            front = str(raw[key]).strip()
            break

    # ── Back ──
    back = None
    for key in ["back", "answer", "response", "a", "back_text", "card_back",
                "back_side", "explanation", "solution"]:
        if key in raw and raw[key]:
            back = str(raw[key]).strip()
            break

    # Must have both front and back
    if not front or not back:
        return None

    # ── Front type ──
    front_type = "question"  # default
    for key in ["front_type", "type", "card_type", "format", "question_type", "kind"]:
        if key in raw and raw[key]:
            val = str(raw[key]).strip().lower()
            if val in ("question", "q"):
                front_type = "question"
            elif val in ("cloze", "fill-in-the-blank", "fill_blank", "blank"):
                front_type = "cloze"
            elif val in ("image", "img", "picture", "visual"):
                front_type = "image"
            else:
                front_type = "question"
            break

    # ── Tags ──
    tags = []
    for key in ["tags", "topics", "categories", "labels", "tag", "topic", "category"]:
        if key in raw:
            val = raw[key]
            if isinstance(val, list):
                tags = [str(t).strip() for t in val if t]
            elif isinstance(val, str):
                # Could be comma-separated or space-separated
                tags = [t.strip() for t in re.split(r'[,;]', val) if t.strip()]
            break
    # Clean tags: lowercase, replace spaces with hyphens
    tags = [t.lower().replace(' ', '-') for t in tags]

    # ── Source quote ──
    source_quote = ""
    for key in ["source_quote", "source", "reference", "quote", "excerpt",
                "original_text", "verbatim", "context"]:
        if key in raw and raw[key]:
            source_quote = str(raw[key]).strip()
            break

    return {
        "front_type": front_type,
        "front": front,
        "back": back,
        "tags": tags,
        "source_quote": source_quote
    }


# ─── Ollama (Local LLM) Integration ────────────────────────────────────────

OLLAMA_URL = "http://localhost:11434"

def check_ollama_available(model: str) -> bool:
    """Check if Ollama is running and the requested model is available."""
    import urllib.request
    import urllib.error

    # Check if Ollama is running
    try:
        req = urllib.request.Request(f"{OLLAMA_URL}/api/tags")
        with urllib.request.urlopen(req, timeout=5) as resp:
            data = json.loads(resp.read().decode('utf-8'))
            available_models = [m["name"] for m in data.get("models", [])]
    except (urllib.error.URLError, ConnectionRefusedError, OSError):
        print("  ERROR: Cannot connect to Ollama at localhost:11434")
        print("  Make sure Ollama is running:")
        print("    brew install ollama     # install (macOS)")
        print("    ollama serve            # start the server")
        return False

    # Check if model is pulled
    # Ollama tags can be "llama3.1:8b" or "llama3.1:latest" — match flexibly
    model_base = model.split(":")[0]
    found = any(model_base in m for m in available_models)

    if not found:
        print(f"  Model '{model}' not found locally. Available: {', '.join(available_models[:5])}")
        print(f"  Pull it with: ollama pull {model}")
        print(f"  Recommended models for this task:")
        print(f"    ollama pull llama3.1:8b      # good balance (4.7 GB)")
        print(f"    ollama pull mistral           # fast, decent quality (4.1 GB)")
        print(f"    ollama pull gemma2:9b          # strong reasoning (5.4 GB)")
        print(f"    ollama pull phi3               # small & fast (2.3 GB)")
        print(f"    ollama pull qwen2.5:7b         # great for structured output (4.4 GB)")
        return False

    print(f"  ✓ Ollama connected, using model: {model}")
    return True


def call_ollama(text: str, has_images: bool, image_names: list[str] = None,
                model: str = "llama3.1:8b", max_cards: int = None,
                low_mem: bool = False, cards_per_page: int = None) -> list[dict]:
    """Call local Ollama to generate flashcards from text."""
    import urllib.request
    import urllib.error

    if not check_ollama_available(model):
        sys.exit(1)

    # Estimate a good card target based on text length
    # ~1 card per 100-200 chars of meaningful content
    estimated_good_count = max(5, len(text) // 150)
    if max_cards:
        target_cards = max_cards
    elif cards_per_page:
        num_pages_est = max(1, len(text) // 2000)
        target_cards = num_pages_est * cards_per_page
    else:
        target_cards = estimated_good_count

    min_cards = max(3, target_cards // 2)

    # Use compact prompt for small models to save context tokens
    is_small_model = any(s in model.lower() for s in [
        'phi3', 'phi-3', 'mini', 'tiny', '1b', '2b', '3b', 'small'
    ])

    user_prompt = build_user_prompt(text, has_images, image_names,
                                     max_cards=target_cards, min_cards=min_cards)

    # Adjust limits based on memory mode
    if low_mem:
        MAX_TEXT_CHARS = 6_000
        ctx_size = 4096
    else:
        MAX_TEXT_CHARS = 24_000
        ctx_size = 32768

    if len(text) > MAX_TEXT_CHARS:
        print(f"  Document is long ({len(text)} chars). Processing in chunks...")
        return _process_ollama_chunks(text, has_images, image_names, model,
                                      max_cards, low_mem, cards_per_page)

    # Pick system prompt: compact for small models / low-mem to save context
    sys_prompt = SYSTEM_PROMPT_COMPACT if (is_small_model or low_mem) else SYSTEM_PROMPT

    full_prompt = f"""{sys_prompt}

{user_prompt}"""

    # Each card is ~150-250 tokens of JSON. Scale output budget to target.
    predict_tokens = min(8192, max(2048, target_cards * 250))

    payload = {
        "model": model,
        "prompt": full_prompt,
        "stream": False,
        "options": {
            "temperature": 0.3,
            "num_predict": predict_tokens,
            "num_ctx": ctx_size,
        },
        "format": "json"
    }

    print(f"  Target: ~{target_cards} cards (min {min_cards})")
    print(f"  Prompt: {'compact' if sys_prompt == SYSTEM_PROMPT_COMPACT else 'full'} ({len(full_prompt)} chars)")
    print(f"  Generating with {model}...")
    
    data = json.dumps(payload).encode('utf-8')
    req = urllib.request.Request(
        f"{OLLAMA_URL}/api/generate",
        data=data,
        headers={"Content-Type": "application/json"}
    )

    try:
        with urllib.request.urlopen(req, timeout=600) as resp:
            result = json.loads(resp.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        error_body = e.read().decode('utf-8') if e.fp else ""
        if "500" in str(e):
            print(f"\n  ❌ Ollama crashed (out of memory).")
            print(f"  Try: --low-mem --model phi3:mini --pages 1-20")
        else:
            print(f"  Ollama error: {e}")
        sys.exit(1)
    except urllib.error.URLError as e:
        print(f"  Ollama error: {e}")
        print(f"  The model may have crashed. Check: ollama ps")
        sys.exit(1)
    except (TimeoutError, socket.timeout):
        print("  Ollama timed out. Try --low-mem, a smaller model, or fewer pages.")
        sys.exit(1)

    response_text = result.get("response", "")

    # Show generation stats
    total_duration = result.get("total_duration", 0)
    if total_duration:
        seconds = total_duration / 1e9
        eval_count = result.get("eval_count", 0)
        tokens_per_sec = eval_count / seconds if seconds > 0 else 0
        print(f"  Generated in {seconds:.1f}s ({tokens_per_sec:.1f} tokens/sec)")

    return _parse_cards_json(response_text)


def _process_ollama_chunks(text: str, has_images: bool, image_names: list[str],
                           model: str, max_cards: int, low_mem: bool = False,
                           cards_per_page: int = None) -> list[dict]:
    """Process long documents in chunks for Ollama's smaller context window."""
    CHUNK_SIZE = 5_000 if low_mem else 20_000
    # Split on paragraph boundaries
    paragraphs = text.split('\n\n')
    chunks = []
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) > CHUNK_SIZE and current_chunk:
            chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk += "\n\n" + para if current_chunk else para
    if current_chunk:
        chunks.append(current_chunk)

    all_cards = []
    cards_per_chunk = (max_cards // len(chunks)) + 1 if max_cards else None

    print(f"  Split into {len(chunks)} chunks ({CHUNK_SIZE} chars each)")

    for i, chunk in enumerate(chunks):
        print(f"\n  Processing chunk {i+1}/{len(chunks)} ({len(chunk)} chars)...")
        cards = call_ollama(chunk, has_images, image_names, model, 
                           cards_per_chunk, low_mem, cards_per_page)
        all_cards.extend(cards)
        print(f"  Running total: {len(all_cards)} cards")

    return all_cards


# ─── Fallback: Rule-based Card Generation (no API needed) ──────────────────

def generate_cards_offline(text: str) -> list[dict]:
    """Generate basic flashcards using rule-based heuristics (no API needed)."""
    print("  Using offline/rule-based card generation...")
    cards = []

    lines = text.split('\n')
    current_topic = "general"

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue

        # Detect headings / topics
        if line.startswith('#'):
            current_topic = line.lstrip('#').strip().lower().replace(' ', '-')
            continue

        # Detect definitions (patterns like "X is Y", "X: Y", "X — Y")
        def_match = re.match(r'^([A-Z][^.!?:]{2,50})(?:\s+is\s+|\s*:\s*|\s*[—–-]\s+)(.{10,})', line)
        if def_match:
            term = def_match.group(1).strip()
            definition = def_match.group(2).strip()
            cards.append({
                "front_type": "question",
                "front": f"What is {term}?",
                "back": line,
                "tags": [current_topic],
                "source_quote": line
            })
            # Cloze version
            cards.append({
                "front_type": "cloze",
                "front": line.replace(term, "{{c1::" + term + "}}"),
                "back": line,
                "tags": [current_topic],
                "source_quote": line
            })
            continue

        # Detect sentences with key facts (longer informational sentences)
        if len(line) > 40 and line[0].isupper() and line.endswith('.'):
            # Find a key term to cloze (first capitalized proper noun or quoted term)
            key_terms = re.findall(r'"([^"]+)"|\'([^\']+)\'|([A-Z][a-z]+(?:\s[A-Z][a-z]+)*)', line)
            if key_terms:
                term = next((t[0] or t[1] or t[2]) for t in key_terms if any(t))
                if term and len(term) > 2:
                    cards.append({
                        "front_type": "cloze",
                        "front": line.replace(term, "{{c1::" + term + "}}", 1),
                        "back": line,
                        "tags": [current_topic],
                        "source_quote": line
                    })

    if not cards:
        # Fallback: create cards from paragraphs
        paragraphs = re.split(r'\n\s*\n', text)
        for para in paragraphs[:50]:
            para = para.strip()
            if len(para) > 30:
                first_sentence = re.split(r'[.!?]', para)[0] + '.'
                cards.append({
                    "front_type": "question",
                    "front": f"What does the text say about: \"{first_sentence[:80]}...\"?",
                    "back": para[:500],
                    "tags": ["general"],
                    "source_quote": para[:500]
                })

    print(f"  Generated {len(cards)} cards (rule-based mode)")
    return cards


# ─── Image Attachment (post-processing) ────────────────────────────────────

# Patterns that indicate a figure/diagram/image reference in text
FIGURE_REF_PATTERNS = [
    r'(?:Fig(?:ure)?|Diagram|Illustration|Chart|Graph|Table|Image|Photo|Exhibit|Plate|Map|Schematic)[\s.]*(\d+[\.\-]?\d*[a-z]?)',
    r'(?:shown|illustrated|depicted|displayed|see|refer to|as in|described in)\s+(?:Fig(?:ure)?|Diagram|the\s+(?:figure|diagram|image|chart))[\s.]*(\d+[\.\-]?\d*[a-z]?)?',
    r'\((?:Fig(?:ure)?|see\s+Fig(?:ure)?)[\s.]*(\d+[\.\-]?\d*[a-z]?)\)',
]

_figure_ref_compiled = [re.compile(p, re.IGNORECASE) for p in FIGURE_REF_PATTERNS]


def detect_figure_reference(text: str) -> bool:
    """Check if text contains a reference to a figure, diagram, chart, etc."""
    for pattern in _figure_ref_compiled:
        if pattern.search(text):
            return True
    return False


def get_figure_number(text: str) -> str:
    """Extract the figure/diagram number from text, e.g. '3.1' from 'Figure 3.1'."""
    for pattern in _figure_ref_compiled:
        m = pattern.search(text)
        if m and m.lastindex and m.group(m.lastindex):
            return m.group(m.lastindex)
    return ""


def extract_page_from_card(card: dict) -> int:
    """Try to determine which page a card's content came from."""
    # Look for [Page N] markers in source_quote or back
    for field in ["source_quote", "back", "front"]:
        text = card.get(field, "")
        m = re.search(r'\[Page\s+(\d+)', text)
        if m:
            return int(m.group(1))
    return 0


def attach_images_to_cards(cards: list[dict], images: list[dict]) -> list[dict]:
    """Post-process cards to attach images from the source document.
    
    Matching strategies (in priority order):
    1. Caption label match: text says "Figure 1-1" and we found a caption
       "Figure 1-1: Single server setup" on a page → exact match
    2. Page proximity with figure ref: text references a figure, match to
       nearest image on the same/adjacent page
    3. Page proximity without ref: card's source page has an image → attach it
    
    Images are appended to the BACK of the card as <img> tags.
    """
    if not images:
        return cards

    # Build mappings
    caption_index = _build_caption_index(images)
    
    page_images = {}
    for img in images:
        pg = img.get("page", 0)
        if pg not in page_images:
            page_images[pg] = []
        page_images[pg].append(img)

    # Log what captions we found
    captioned = [img for img in images if img.get("caption_label")]
    if captioned:
        print(f"  Found {len(captioned)} captioned images:")
        for img in captioned[:8]:
            print(f"    {img['name']} → \"{img['caption']}\"")
        if len(captioned) > 8:
            print(f"    ... and {len(captioned) - 8} more")
    
    attached_count = 0
    caption_match_count = 0
    fig_ref_count = 0

    for card in cards:
        front = card.get("front", "")
        back = card.get("back", "")
        source = card.get("source_quote", "")
        combined_text = f"{front} {back} {source}"

        # Skip if this card already has an <img> tag
        if '<img' in back or '<img' in front:
            continue

        card_page = extract_page_from_card(card)
        best_img_name = None
        is_fig_ref = detect_figure_reference(combined_text)

        if is_fig_ref:
            # Strategy 1: Match by caption label (highest confidence)
            fig_num = get_figure_number(combined_text)
            if fig_num and fig_num in caption_index:
                best_img_name = caption_index[fig_num]
                caption_match_count += 1
            
            # Strategy 2: Match by page proximity
            if not best_img_name and card_page > 0:
                for offset in [0, -1, 1, -2, 2]:
                    target_page = card_page + offset
                    if target_page in page_images and page_images[target_page]:
                        best_img_name = page_images[target_page][0]["name"]
                        break
            
            # Strategy 3: Figure number as global image index (last resort)
            if not best_img_name and fig_num:
                try:
                    idx = int(fig_num.split('.')[0].split('-')[0]) - 1
                    if 0 <= idx < len(images):
                        best_img_name = images[idx]["name"]
                except ValueError:
                    pass
            
            if best_img_name:
                fig_ref_count += 1

        elif card_page > 0:
            # Strategy 4: No figure ref, but card is from a page with images
            if card_page in page_images and page_images[card_page]:
                best_img_name = page_images[card_page][0]["name"]

        if best_img_name:
            card["back"] = card["back"] + f'<br><br><img src="{best_img_name}">'
            attached_count += 1

    if attached_count > 0:
        parts = []
        if caption_match_count:
            parts.append(f"{caption_match_count} by caption label")
        if fig_ref_count - caption_match_count > 0:
            parts.append(f"{fig_ref_count - caption_match_count} by page+figure ref")
        if attached_count - fig_ref_count > 0:
            parts.append(f"{attached_count - fig_ref_count} by page proximity")
        print(f"  Attached images to {attached_count} cards ({', '.join(parts)})")
    else:
        print(f"  No image-text matches found (images still saved to media folder)")

    return cards


def save_used_images(images: list[dict], cards: list[dict], image_dir: str):
    """Save images to a folder for copying into Anki's collection.media.
    
    Saves images that are referenced in cards via <img> tags.
    Also saves ALL extracted images if any figure references were detected,
    since the LLM may have rephrased text and dropped the explicit reference.
    """
    # Collect all image filenames referenced in card HTML
    referenced = set()
    for card in cards:
        for field in ["front", "back"]:
            text = card.get(field, "")
            refs = re.findall(r'<img\s+src="([^"]+)"', text)
            referenced.update(refs)

    # If we have referenced images, save those.
    # If not but we DO have images from the document, save all of them
    # (the user may want to manually link them or the references may have 
    # been lost during LLM processing).
    images_to_save = []
    if referenced:
        images_to_save = [img for img in images if img["name"] in referenced]
    
    # Always save all images — they came from the document and the user
    # needs them in collection.media for any image cards to work
    if not images_to_save:
        images_to_save = images

    if not images_to_save:
        return False

    os.makedirs(image_dir, exist_ok=True)
    saved = 0
    for img in images_to_save:
        img_path = os.path.join(image_dir, img["name"])
        with open(img_path, 'wb') as f:
            f.write(base64.b64decode(img["data"]))
        saved += 1

    if saved > 0:
        in_cards = len(referenced)
        print(f"  Saved {saved} images to: {image_dir}")
        if in_cards:
            print(f"    ({in_cards} directly referenced in cards)")
        print(f"  ⚠️  Copy ALL files from this folder to your Anki media folder:")
        print(f"    macOS:   ~/Library/Application Support/Anki2/<profile>/collection.media/")
        print(f"    Linux:   ~/.local/share/Anki2/<profile>/collection.media/")
        print(f"    Windows: %APPDATA%\\Anki2\\<profile>\\collection.media\\")
    return saved > 0

def export_to_csv(cards: list[dict], output_path: str, images: list[dict] = None,
                  image_dir: str = None):
    """
    Export cards to Anki-importable CSV.
    Format: Front, Back, Tags
    Anki import settings: separator=tab, allow HTML, field 1=Front, field 2=Back, field 3=Tags
    """
    # Build image lookup: map base names (without extension) to full filenames
    image_lookup = {}
    if images:
        for img in images:
            full_name = img["name"]                          # e.g. "img_15_9.png"
            base_name = full_name.rsplit('.', 1)[0]          # e.g. "img_15_9"
            image_lookup[base_name] = full_name
            image_lookup[full_name] = full_name              # also map full name to itself

    # Save only images that are actually used in cards
    if images and image_dir:
        save_used_images(images, cards, image_dir)

    def resolve_image_ref(text: str) -> str:
        """Find image references in text and return the correct full filename."""
        # Try patterns: img_15_9.png, img_15_9, img_15
        for pattern in [r'(img_[\w]+\.(?:png|jpg|jpeg|gif|bmp|webp))',  # with extension
                        r'(img_[\w]+)']:                                 # without extension
            match = re.search(pattern, text)
            if match:
                ref = match.group(1)
                # Direct match
                if ref in image_lookup:
                    return image_lookup[ref]
                # Strip extension and try base name
                base = ref.rsplit('.', 1)[0]
                if base in image_lookup:
                    return image_lookup[base]
                # Fuzzy: find any image whose name contains this reference
                for key, val in image_lookup.items():
                    if ref in key or key in ref:
                        return val
        return None

    with open(output_path, 'w', encoding='utf-8') as f:
        # Write metadata comment for Anki import
        f.write("#separator:tab\n")
        f.write("#html:true\n")
        f.write("#tags column:3\n")

        for card in cards:
            front = card["front"]
            back = card["back"]

            # Format front based on type
            if card.get("front_type") == "cloze":
                # Convert cloze syntax for display (Anki Basic doesn't have native cloze,
                # so we render it as a fill-in-the-blank with underscores)
                display_front = re.sub(r'\{\{c\d+::([^}]+)\}\}', r'_____', front)
                # Put the cloze answer hint in the back
                cloze_answers = re.findall(r'\{\{c\d+::([^}]+)\}\}', front)
                if cloze_answers:
                    cloze_context = re.sub(r'\{\{c\d+::([^}]+)\}\}', r'<b>\1</b>', front)
                    back = f"{cloze_context}<br><br><i>Source:</i> {back}"
                front = display_front

            elif card.get("front_type") == "image":
                # Resolve image reference to actual filename
                img_filename = resolve_image_ref(front)
                if img_filename:
                    img_tag = f'<img src="{img_filename}">'
                    # Remove the raw image reference from the question text
                    clean_front = re.sub(r'img_[\w]+(?:\.(?:png|jpg|jpeg|gif|bmp|webp))?', '', front).strip()
                    clean_front = re.sub(r'^[\s,.:;-]+', '', clean_front).strip()  # clean leading punctuation
                    front = f"{img_tag}<br>{clean_front}" if clean_front else img_tag

            # Also resolve any RAW image references in the back (e.g. "img_6_0.png" as plain text)
            # But SKIP if the back already has proper <img> tags (from attach_images_to_cards)
            if images and '<img' not in back:
                img_in_back = resolve_image_ref(back)
                if img_in_back:
                    back = re.sub(
                        r'img_[\w]+(?:\.(?:png|jpg|jpeg|gif|bmp|webp))?',
                        f'<img src="{img_in_back}">',
                        back
                    )

            # Format back with source attribution
            if card.get("source_quote") and card["source_quote"] != back:
                back_html = f"{back}<br><br><small><i>Source: \"{card['source_quote'][:200]}\"</i></small>"
            else:
                back_html = back

            # Format tags (space-separated for Anki)
            tags = " ".join(card.get("tags", []))

            # Clean fields: replace any tabs/newlines that would break the format
            front = front.replace('\t', ' ').replace('\n', '<br>')
            back_html = back_html.replace('\t', ' ').replace('\n', '<br>')
            tags = tags.replace('\t', ' ').replace('\n', ' ')

            # Write tab-separated line (no quoting — Anki handles HTML natively)
            f.write(f"{front}\t{back_html}\t{tags}\n")

    card_count = len(cards)
    print(f"\n  ✓ Exported {card_count} cards to: {output_path}")
    print(f"  Import into Anki: File → Import → select the CSV file")
    print(f"  Settings: Type=Basic, Separator=Tab, Allow HTML=Yes")


# ─── AnkiConnect Integration ───────────────────────────────────────────────

def push_to_anki_connect(cards: list[dict], deck_name: str, images: list[dict] = None):
    """Push cards to Anki via AnkiConnect (requires Anki running with AnkiConnect plugin)."""
    import urllib.request
    import urllib.error

    ANKI_CONNECT_URL = "http://localhost:8765"

    def anki_request(action: str, **params):
        payload = json.dumps({"action": action, "version": 6, "params": params})
        req = urllib.request.Request(ANKI_CONNECT_URL, data=payload.encode('utf-8'),
                                     headers={"Content-Type": "application/json"})
        try:
            with urllib.request.urlopen(req, timeout=10) as resp:
                result = json.loads(resp.read().decode('utf-8'))
                if result.get("error"):
                    print(f"  AnkiConnect error: {result['error']}")
                return result.get("result")
        except urllib.error.URLError:
            print("  ERROR: Cannot connect to AnkiConnect.")
            print("  Make sure Anki is running and AnkiConnect plugin is installed.")
            print("  Plugin code: 2055492159")
            return None

    # Check connection
    version = anki_request("version")
    if version is None:
        return False

    print(f"  Connected to AnkiConnect v{version}")

    # Create deck
    anki_request("createDeck", deck=deck_name)
    print(f"  Created/verified deck: {deck_name}")

    # Store images first
    if images:
        for img in images:
            anki_request("storeMediaFile",
                        filename=img["name"],
                        data=img["data"])
        print(f"  Uploaded {len(images)} images")

    # Build image lookup for resolving references
    image_lookup = {}
    if images:
        for img in images:
            full_name = img["name"]
            base_name = full_name.rsplit('.', 1)[0]
            image_lookup[base_name] = full_name
            image_lookup[full_name] = full_name

    def resolve_img(text):
        for pat in [r'(img_[\w]+\.(?:png|jpg|jpeg|gif|bmp|webp))', r'(img_[\w]+)']:
            m = re.search(pat, text)
            if m:
                ref = m.group(1)
                if ref in image_lookup: return image_lookup[ref]
                base = ref.rsplit('.', 1)[0]
                if base in image_lookup: return image_lookup[base]
        return None

    # Add notes
    notes = []
    for card in cards:
        front = card["front"]
        back = card["back"]

        # Format for Anki
        if card.get("front_type") == "cloze":
            display_front = re.sub(r'\{\{c\d+::([^}]+)\}\}', r'_____', front)
            cloze_context = re.sub(r'\{\{c\d+::([^}]+)\}\}', r'<b>\1</b>', front)
            back = f"{cloze_context}<br><br><i>Source:</i> {back}"
            front = display_front
        elif card.get("front_type") == "image":
            img_filename = resolve_img(front)
            if img_filename:
                clean_front = re.sub(r'img_[\w]+(?:\.(?:png|jpg|jpeg|gif|bmp|webp))?', '', front).strip()
                clean_front = re.sub(r'^[\s,.:;-]+', '', clean_front).strip()
                front = f'<img src="{img_filename}"><br>{clean_front}' if clean_front else f'<img src="{img_filename}">'

        note = {
            "deckName": deck_name,
            "modelName": "Basic",
            "fields": {
                "Front": front,
                "Back": back
            },
            "tags": card.get("tags", []),
            "options": {
                "allowDuplicate": False,
                "duplicateScope": "deck"
            }
        }
        notes.append(note)

    # Add notes in batch
    result = anki_request("addNotes", notes=notes)
    if result:
        added = sum(1 for r in result if r is not None)
        dupes = sum(1 for r in result if r is None)
        print(f"  ✓ Added {added} cards to deck '{deck_name}'")
        if dupes:
            print(f"  ↳ {dupes} duplicates skipped")
    return True


# ─── Main ───────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate Anki flashcard decks from PDF, DOCX, or TXT files",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python anki_generator.py textbook.pdf
  python anki_generator.py notes.docx -o my_cards.csv -d "Bio 101"
  python anki_generator.py lecture.txt --anki-connect --deck "Physics"
  python anki_generator.py study.pdf --offline  # No LLM needed

  # Local LLM via Ollama (free, private, no API key):
  python anki_generator.py textbook.pdf --ollama
  python anki_generator.py notes.docx --ollama --model mistral
  python anki_generator.py lecture.txt --ollama --model gemma2:9b -m 30

  # For 8GB Macs (269-page PDF etc.):
  python anki_generator.py big.pdf --ollama --low-mem --pages 1-30
  python anki_generator.py big.pdf --ollama --low-mem --model phi3:mini
        """
    )
    parser.add_argument("input_file", help="Path to input file (PDF, DOCX, or TXT)")
    parser.add_argument("--output", "-o", help="Output CSV filename")
    parser.add_argument("--deck", "-d", help="Anki deck name (default: from filename)")
    parser.add_argument("--anki-connect", action="store_true",
                        help="Push cards directly to Anki via AnkiConnect")
    parser.add_argument("--api-key", help="Anthropic API key (or set ANTHROPIC_API_KEY)")
    parser.add_argument("--ollama", action="store_true",
                        help="Use local Ollama LLM instead of Claude API (free, no key needed)")
    parser.add_argument("--model", default="llama3.1:8b",
                        help="Ollama model to use (default: llama3.1:8b)")
    parser.add_argument("--max-cards", "-m", type=int, help="Target number of cards to generate")
    parser.add_argument("--cards-per-page", "-cpp", type=int, default=None,
                        help="Target cards per page (e.g. 5 = ~5 cards per page of content)")
    parser.add_argument("--pages", "-p",
                        help="Page range for PDFs, e.g. '1-20' or '50-75' (default: all)")
    parser.add_argument("--low-mem", action="store_true",
                        help="Low memory mode for 8GB Macs (smaller chunks, smaller context)")
    parser.add_argument("--offline", action="store_true",
                        help="Use rule-based generation (no LLM at all)")
    parser.add_argument("--verbose", "-v", action="store_true", help="Show detailed progress")

    args = parser.parse_args()

    # Validate input file
    input_path = Path(args.input_file)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)

    ext = input_path.suffix.lower()
    if ext not in ('.pdf', '.docx', '.doc', '.txt', '.md'):
        print(f"ERROR: Unsupported file type: {ext}")
        print("Supported: .pdf, .docx, .txt, .md")
        sys.exit(1)

    # Set defaults
    deck_name = args.deck or input_path.stem.replace('_', ' ').replace('-', ' ').title()
    output_file = args.output or f"{input_path.stem}_anki_deck.csv"

    # Parse page range if given
    page_range = None
    if args.pages:
        try:
            parts = args.pages.split('-')
            page_range = (int(parts[0]), int(parts[1]))
            if page_range[0] < 1 or page_range[1] < page_range[0]:
                raise ValueError
        except (ValueError, IndexError):
            print(f"ERROR: Invalid page range '{args.pages}'. Use format: 1-20")
            sys.exit(1)

    # Determine backend
    if args.offline:
        backend = "offline (rule-based)"
    elif args.ollama:
        backend = f"Ollama ({args.model})"
        if args.low_mem:
            backend += " [low-mem]"
    else:
        backend = "Claude API"

    print(f"╔══════════════════════════════════════════════╗")
    print(f"║        Anki Deck Generator                   ║")
    print(f"╚══════════════════════════════════════════════╝")
    print(f"  Input:   {input_path.name}")
    print(f"  Deck:    {deck_name}")
    print(f"  Output:  {output_file}")
    print(f"  Backend: {backend}")
    if page_range:
        print(f"  Pages:   {page_range[0]}-{page_range[1]}")
    print()

    # Step 1: Read the document
    print("📄 Step 1: Reading document...")
    if ext == '.pdf' and not POPPLER_AVAILABLE:
        print("  ⚠️  poppler not found — image extraction will be limited.")
        print("     Install for best results:")
        print("       macOS:  brew install poppler")
        print("       Linux:  sudo apt install poppler-utils")
        print()
    if ext == '.pdf':
        text, images = read_pdf(str(input_path), page_range=page_range)
    elif ext in ('.docx', '.doc'):
        text, images = read_docx(str(input_path))
    else:
        text, images = read_txt(str(input_path))

    if not text.strip():
        print("ERROR: No text could be extracted from the file.")
        print("If this is a scanned PDF, OCR support may be needed.")
        sys.exit(1)

    print(f"  Extracted {len(text)} characters of text")
    if images:
        print(f"  Found {len(images)} embedded images")

    # Filter out boilerplate (TOC, index, acknowledgements, etc.)
    original_len = len(text)
    text = filter_content(text)
    filtered_len = len(text)
    if original_len != filtered_len:
        removed_pct = ((original_len - filtered_len) / original_len) * 100
        print(f"  After filtering: {filtered_len} chars ({removed_pct:.0f}% boilerplate removed)")

    if not text.strip():
        print("ERROR: All content was filtered out. Try with --pages to target specific pages.")
        sys.exit(1)

    print()

    # Step 2: Generate flashcards
    print("🧠 Step 2: Generating flashcards...")
    if args.offline:
        cards = generate_cards_offline(text)
    elif args.ollama:
        image_names = [img["name"] for img in images] if images else []
        cards = call_ollama(
            text,
            has_images=bool(images),
            image_names=image_names,
            model=args.model,
            max_cards=args.max_cards,
            low_mem=args.low_mem,
            cards_per_page=args.cards_per_page
        )
    else:
        image_names = [img["name"] for img in images] if images else []
        cards = call_claude_api(
            text,
            has_images=bool(images),
            image_names=image_names,
            api_key=args.api_key,
            max_cards=args.max_cards
        )

    if not cards:
        print("ERROR: No flashcards were generated.")
        sys.exit(1)

    print(f"  Generated {len(cards)} flashcards")

    # Step 2b: Attach images to cards that reference figures/diagrams
    if images:
        print("\n🖼️  Step 2b: Attaching images to figure references...")
        cards = attach_images_to_cards(cards, images)

    # Show summary of card types
    type_counts = {}
    for c in cards:
        t = c.get("front_type", "question")
        type_counts[t] = type_counts.get(t, 0) + 1
    print(f"  Types: {', '.join(f'{t}={n}' for t, n in type_counts.items())}")

    # Collect all unique tags
    all_tags = set()
    for c in cards:
        all_tags.update(c.get("tags", []))
    print(f"  Topics: {', '.join(sorted(all_tags)[:10])}")
    if len(all_tags) > 10:
        print(f"          ... and {len(all_tags) - 10} more")
    print()

    # Step 3: Export
    print("📦 Step 3: Exporting...")

    # Always export CSV
    image_dir = None
    if images:
        image_dir = str(Path(output_file).with_suffix('')) + "_media"
    export_to_csv(cards, output_file, images, image_dir)

    # Optionally push to AnkiConnect
    if args.anki_connect:
        print("\n🔗 Pushing to AnkiConnect...")
        push_to_anki_connect(cards, deck_name, images)

    print("\n✅ Done!")

    if args.verbose:
        print("\n─── Sample Cards ───")
        for card in cards[:3]:
            print(f"\n  Front [{card.get('front_type', '?')}]: {card['front'][:100]}")
            print(f"  Back:  {card['back'][:100]}")
            print(f"  Tags:  {card.get('tags', [])}")


if __name__ == "__main__":
    main()
